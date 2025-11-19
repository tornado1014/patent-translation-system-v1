#!/usr/bin/env python3
"""
특허 번역 자동화 시스템 - GUI 버전
PyQt6 기반 데스크톱 애플리케이션
"""

import sys
import os
from pathlib import Path
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QLineEdit, QTextEdit, QComboBox, QFileDialog,
    QProgressBar, QCheckBox, QGroupBox, QMessageBox, QTabWidget,
    QTableWidget, QTableWidgetItem, QListWidget, QListWidgetItem, QSpinBox
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtGui import QFont, QIcon

# src 디렉토리를 Python 경로에 추가
sys.path.insert(0, str(Path(__file__).parent / "src"))

from pipeline import TranslationPipeline
from tm_manager import TranslationMemory
from section_parser import PatentSectionParser


class TranslationThread(QThread):
    """번역 작업을 백그라운드에서 실행하는 스레드"""

    progress = pyqtSignal(str)  # 진행 상황 메시지
    finished = pyqtSignal(dict)  # 완료 시 결과
    error = pyqtSignal(str)  # 오류 발생 시

    def __init__(self, input_file, output_file, doc_type, use_review, save_tm, auto_section=False):
        super().__init__()
        self.input_file = input_file
        self.output_file = output_file
        self.doc_type = doc_type
        self.use_review = use_review
        self.save_tm = save_tm
        self.auto_section = auto_section

    def run(self):
        try:
            self.progress.emit("🚀 번역 파이프라인 초기화 중...")
            pipeline = TranslationPipeline()

            # 입력 파일 읽기
            self.progress.emit(f"📄 파일 읽기: {self.input_file}")
            input_path = Path(self.input_file)
            file_ext = input_path.suffix.lower()

            if file_ext == '.txt':
                with open(self.input_file, 'r', encoding='utf-8') as f:
                    source_text = f.read()
            elif file_ext == '.docx':
                from docx import Document
                doc = Document(self.input_file)
                source_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                self.progress.emit(f"📄 Word 파일 읽기 완료: {len(doc.paragraphs)}개 문단")
            elif file_ext == '.pdf':
                from docling.document_converter import DocumentConverter
                converter = DocumentConverter()
                result = converter.convert(self.input_file)
                source_text = result.document.export_to_markdown()
                self.progress.emit("📄 PDF 파일 변환 완료")
            else:
                raise ValueError(f"지원하지 않는 파일 형식: {file_ext}")

            # 자동 섹션 분류 모드
            if self.auto_section:
                self.progress.emit("🤖 자동 섹션 분류 시작...")

                parser = PatentSectionParser()
                sections = parser.parse_document(source_text)

                # 섹션별 번역 결과 저장
                translated_sections = {}
                total_sections = sum(len(items) for items in sections.values())
                current = 0

                for section_type, section_list in sections.items():
                    if not section_list:
                        continue

                    translated_sections[section_type] = []

                    for i, section in enumerate(section_list, 1):
                        current += 1
                        doc_type = parser.get_document_type_from_section(section.section_type)

                        self.progress.emit(
                            f"📝 번역 중 ({current}/{total_sections}): "
                            f"{section_type.upper()} #{i} - {doc_type}"
                        )

                        # 섹션별 번역
                        result = pipeline.translate_document(
                            source_text=section.content,
                            document_type=doc_type,
                            use_self_review=self.use_review,
                            save_to_tm=self.save_tm
                        )

                        if result["success"]:
                            translated_sections[section_type].append(
                                (section, result["translation"])
                            )
                        else:
                            raise Exception(f"섹션 번역 실패: {result.get('error')}")

                # 번역된 섹션 재구성
                self.progress.emit("🔄 번역 문서 재구성 중...")
                translation = parser.reconstruct_document(translated_sections)

                # 결과 생성
                result = {
                    "success": True,
                    "translation": translation,
                    "sections": {k: len(v) for k, v in sections.items()},
                    "auto_section": True
                }

            # 일반 번역 모드
            else:
                self.progress.emit("🔄 번역 시작...")
                result = pipeline.translate_document(
                    source_text=source_text,
                    document_type=self.doc_type,
                    use_self_review=self.use_review,
                    save_to_tm=self.save_tm
                )

            if result["success"]:
                translation = result["translation"]

                # 출력 파일 저장
                self.progress.emit(f"💾 저장 중: {self.output_file}")
                output_path = Path(self.output_file)
                output_path.parent.mkdir(parents=True, exist_ok=True)
                output_ext = output_path.suffix.lower()

                if output_ext == '.docx':
                    from docx import Document
                    doc = Document()
                    for line in translation.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line)
                    doc.save(str(output_path))
                else:
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(translation)

                # QA 리포트 저장
                if "qa_result" in result:
                    qa_report = pipeline.qa_checker.generate_report(result["qa_result"])
                    qa_path = output_path.with_suffix('.qa.txt')
                    with open(qa_path, 'w', encoding='utf-8') as f:
                        f.write(qa_report)

                self.progress.emit("✅ 번역 완료!")
                pipeline.close()
                self.finished.emit(result)
            else:
                raise Exception(result.get('error', '알 수 없는 오류'))

        except Exception as e:
            self.error.emit(str(e))


class PatentTranslatorGUI(QMainWindow):
    """특허 번역 GUI 메인 윈도우"""

    def __init__(self):
        super().__init__()
        self.translation_thread = None
        self.init_ui()

    def init_ui(self):
        """UI 초기화"""
        self.setWindowTitle("특허 번역 자동화 시스템 v1.0")
        self.setGeometry(100, 100, 900, 700)

        # 중앙 위젯
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        # 메인 레이아웃
        main_layout = QVBoxLayout()
        central_widget.setLayout(main_layout)

        # 탭 위젯
        tabs = QTabWidget()
        main_layout.addWidget(tabs)

        # 번역 탭
        translate_tab = self.create_translate_tab()
        tabs.addTab(translate_tab, "📝 번역")

        # TM 통계 탭
        tm_tab = self.create_tm_tab()
        tabs.addTab(tm_tab, "📊 Translation Memory")

        # 로그 뷰어 탭
        log_tab = self.create_log_viewer_tab()
        tabs.addTab(log_tab, "📄 로그")

        # 설정 탭
        settings_tab = self.create_settings_tab()
        tabs.addTab(settings_tab, "⚙️ 설정")

        # 상태바
        self.statusBar().showMessage("준비 완료")

    def create_translate_tab(self):
        """번역 탭 생성"""
        tab = QWidget()
        layout = QVBoxLayout()
        tab.setLayout(layout)

        # 타이틀
        title = QLabel("🚀 특허 번역 자동화 시스템")
        title_font = QFont()
        title_font.setPointSize(16)
        title_font.setBold(True)
        title.setFont(title_font)
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)

        # 입력 파일 그룹
        input_group = QGroupBox("📥 입력 파일")
        input_layout = QHBoxLayout()
        input_group.setLayout(input_layout)

        self.input_file_edit = QLineEdit()
        self.input_file_edit.setPlaceholderText("입력 파일 선택 (.txt, .docx, .pdf)")
        input_layout.addWidget(self.input_file_edit)

        input_btn = QPushButton("파일 선택")
        input_btn.clicked.connect(self.select_input_file)
        input_layout.addWidget(input_btn)

        layout.addWidget(input_group)

        # 출력 파일 그룹
        output_group = QGroupBox("📤 출력 파일")
        output_layout = QHBoxLayout()
        output_group.setLayout(output_layout)

        self.output_file_edit = QLineEdit()
        self.output_file_edit.setPlaceholderText("출력 파일 선택 (.txt, .docx)")
        output_layout.addWidget(self.output_file_edit)

        output_btn = QPushButton("파일 선택")
        output_btn.clicked.connect(self.select_output_file)
        output_layout.addWidget(output_btn)

        layout.addWidget(output_group)

        # 옵션 그룹
        options_group = QGroupBox("⚙️ 번역 옵션")
        options_layout = QVBoxLayout()
        options_group.setLayout(options_layout)

        # 자동 섹션 분류 체크박스
        self.auto_section_checkbox = QCheckBox("🤖 자동 섹션 분류 (전체 명세서 입력 시)")
        self.auto_section_checkbox.setChecked(False)
        self.auto_section_checkbox.toggled.connect(self.toggle_auto_section)
        options_layout.addWidget(self.auto_section_checkbox)

        # 문서 유형 (자동 분류 시 비활성화)
        doc_type_layout = QHBoxLayout()
        doc_type_layout.addWidget(QLabel("문서 유형:"))
        self.doc_type_combo = QComboBox()
        self.doc_type_combo.addItems(["청구항 (claim)", "명세서 (specification)", "요약서 (abstract)"])
        doc_type_layout.addWidget(self.doc_type_combo)
        doc_type_layout.addStretch()
        options_layout.addLayout(doc_type_layout)

        # 체크박스 옵션
        self.review_checkbox = QCheckBox("자체 검수 활성화 (권장)")
        self.review_checkbox.setChecked(True)
        options_layout.addWidget(self.review_checkbox)

        self.tm_checkbox = QCheckBox("Translation Memory에 저장")
        self.tm_checkbox.setChecked(True)
        options_layout.addWidget(self.tm_checkbox)

        layout.addWidget(options_group)

        # 번역 버튼
        self.translate_btn = QPushButton("🚀 번역 시작")
        self.translate_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                font-size: 14pt;
                font-weight: bold;
                padding: 10px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QPushButton:disabled {
                background-color: #cccccc;
            }
        """)
        self.translate_btn.clicked.connect(self.start_translation)
        layout.addWidget(self.translate_btn)

        # 진행 상황
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 0)  # Indeterminate mode
        self.progress_bar.hide()
        layout.addWidget(self.progress_bar)

        # 로그
        log_label = QLabel("📋 진행 상황:")
        layout.addWidget(log_label)

        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(200)
        layout.addWidget(self.log_text)

        return tab

    def create_tm_tab(self):
        """TM 검색 및 통계 탭 생성"""
        tab = QWidget()
        layout = QVBoxLayout()
        tab.setLayout(layout)

        # 타이틀
        title = QLabel("📊 Translation Memory")
        title_font = QFont()
        title_font.setPointSize(14)
        title_font.setBold(True)
        title.setFont(title_font)
        layout.addWidget(title)

        # 검색 섹션
        search_group = QGroupBox("🔍 TM 검색")
        search_layout = QVBoxLayout()
        search_group.setLayout(search_layout)

        # 검색 입력
        search_input_layout = QHBoxLayout()
        search_input_layout.addWidget(QLabel("검색어:"))
        self.tm_search_input = QLineEdit()
        self.tm_search_input.setPlaceholderText("영문 또는 한글 텍스트 입력...")
        self.tm_search_input.returnPressed.connect(self.search_tm)
        search_input_layout.addWidget(self.tm_search_input)
        search_layout.addLayout(search_input_layout)

        # 유사도 임계값 설정
        threshold_layout = QHBoxLayout()
        threshold_layout.addWidget(QLabel("유사도 임계값:"))
        self.tm_threshold_spinbox = QSpinBox()
        self.tm_threshold_spinbox.setRange(50, 100)
        self.tm_threshold_spinbox.setValue(70)
        self.tm_threshold_spinbox.setSuffix("%")
        threshold_layout.addWidget(self.tm_threshold_spinbox)
        threshold_layout.addStretch()
        search_layout.addLayout(threshold_layout)

        # 검색 버튼
        search_btn = QPushButton("🔍 검색")
        search_btn.clicked.connect(self.search_tm)
        search_layout.addWidget(search_btn)

        # 검색 결과
        self.tm_search_results = QTextEdit()
        self.tm_search_results.setReadOnly(True)
        self.tm_search_results.setPlaceholderText("검색 결과가 여기에 표시됩니다...")
        search_layout.addWidget(self.tm_search_results)

        layout.addWidget(search_group)

        # 통계 섹션
        stats_group = QGroupBox("📊 TM 통계")
        stats_layout = QVBoxLayout()
        stats_group.setLayout(stats_layout)

        # 새로고침 버튼
        refresh_btn = QPushButton("🔄 새로고침")
        refresh_btn.clicked.connect(self.refresh_tm_stats)
        stats_layout.addWidget(refresh_btn)

        # 통계 텍스트
        self.tm_stats_text = QTextEdit()
        self.tm_stats_text.setReadOnly(True)
        self.tm_stats_text.setMaximumHeight(150)
        stats_layout.addWidget(self.tm_stats_text)

        layout.addWidget(stats_group)

        # 초기 로드
        self.refresh_tm_stats()

        return tab

    def create_log_viewer_tab(self):
        """로그 뷰어 탭 생성"""
        tab = QWidget()
        layout = QVBoxLayout()
        tab.setLayout(layout)

        # 타이틀
        title = QLabel("📄 번역 로그 뷰어")
        title_font = QFont()
        title_font.setPointSize(14)
        title_font.setBold(True)
        title.setFont(title_font)
        layout.addWidget(title)

        # 로그 파일 목록
        log_list_group = QGroupBox("📂 로그 파일 목록")
        log_list_layout = QVBoxLayout()
        log_list_group.setLayout(log_list_layout)

        # 새로고침 버튼
        refresh_logs_btn = QPushButton("🔄 새로고침")
        refresh_logs_btn.clicked.connect(self.refresh_log_list)
        log_list_layout.addWidget(refresh_logs_btn)

        # 로그 파일 리스트
        self.log_file_list = QListWidget()
        self.log_file_list.itemClicked.connect(self.load_log_file)
        log_list_layout.addWidget(self.log_file_list)

        layout.addWidget(log_list_group)

        # 로그 내용 뷰어
        log_content_group = QGroupBox("📋 로그 내용")
        log_content_layout = QVBoxLayout()
        log_content_group.setLayout(log_content_layout)

        # 로그 필터
        filter_layout = QHBoxLayout()
        filter_layout.addWidget(QLabel("필터:"))
        self.log_filter_combo = QComboBox()
        self.log_filter_combo.addItems(["전체", "INFO", "WARNING", "ERROR", "DEBUG"])
        self.log_filter_combo.currentTextChanged.connect(self.filter_log_content)
        filter_layout.addWidget(self.log_filter_combo)
        filter_layout.addStretch()
        log_content_layout.addLayout(filter_layout)

        # 로그 텍스트
        self.log_content_text = QTextEdit()
        self.log_content_text.setReadOnly(True)
        self.log_content_text.setFont(QFont("Courier", 10))
        log_content_layout.addWidget(self.log_content_text)

        # 로그 지우기 버튼
        clear_log_btn = QPushButton("🗑️ 표시된 로그 지우기")
        clear_log_btn.clicked.connect(lambda: self.log_content_text.clear())
        log_content_layout.addWidget(clear_log_btn)

        layout.addWidget(log_content_group)

        # 초기 로그 목록 로드
        self.refresh_log_list()

        return tab

    def create_settings_tab(self):
        """설정 탭 생성"""
        tab = QWidget()
        layout = QVBoxLayout()
        tab.setLayout(layout)

        # 타이틀
        title = QLabel("⚙️ 시스템 설정")
        title_font = QFont()
        title_font.setPointSize(14)
        title_font.setBold(True)
        title.setFont(title_font)
        layout.addWidget(title)

        # 탭 위젯 생성 (설정 항목별)
        settings_tabs = QTabWidget()
        layout.addWidget(settings_tabs)

        # API 설정 탭
        api_tab = self.create_api_settings_tab()
        settings_tabs.addTab(api_tab, "🔑 API 설정")

        # 용어집 편집 탭
        terminology_tab = self.create_terminology_tab()
        settings_tabs.addTab(terminology_tab, "📚 용어집")

        # 시스템 정보 탭
        info_tab = self.create_info_tab()
        settings_tabs.addTab(info_tab, "ℹ️ 정보")

        return tab

    def create_api_settings_tab(self):
        """API 설정 탭"""
        tab = QWidget()
        layout = QVBoxLayout()
        tab.setLayout(layout)

        # API 키 설정
        api_group = QGroupBox("🔑 Claude API 설정")
        api_layout = QVBoxLayout()
        api_group.setLayout(api_layout)

        # API 키 입력
        key_layout = QHBoxLayout()
        key_layout.addWidget(QLabel("API 키:"))
        self.api_key_edit = QLineEdit()
        self.api_key_edit.setEchoMode(QLineEdit.EchoMode.Password)
        self.api_key_edit.setPlaceholderText("sk-ant-...")

        # 현재 설정된 API 키 로드
        try:
            import os
            current_key = os.getenv("ANTHROPIC_API_KEY", "")
            if current_key:
                self.api_key_edit.setText(current_key)
        except:
            pass

        key_layout.addWidget(self.api_key_edit)
        api_layout.addLayout(key_layout)

        # API 키 표시/숨기기 버튼
        show_key_btn = QPushButton("👁️ 표시/숨기기")
        show_key_btn.clicked.connect(self.toggle_api_key_visibility)
        api_layout.addWidget(show_key_btn)

        # API 키 저장 버튼
        save_key_btn = QPushButton("💾 API 키 저장")
        save_key_btn.clicked.connect(self.save_api_key)
        api_layout.addWidget(save_key_btn)

        layout.addWidget(api_group)

        # 파일 경로 설정
        paths_group = QGroupBox("📁 파일 경로")
        paths_layout = QVBoxLayout()
        paths_group.setLayout(paths_layout)

        # 스타일 가이드 경로
        style_guide_layout = QHBoxLayout()
        style_guide_layout.addWidget(QLabel("스타일 가이드:"))
        self.style_guide_path_edit = QLineEdit()
        self.style_guide_path_edit.setText("config/style_guide.json")
        self.style_guide_path_edit.setReadOnly(True)
        style_guide_layout.addWidget(self.style_guide_path_edit)
        browse_style_btn = QPushButton("찾아보기")
        browse_style_btn.clicked.connect(lambda: self.browse_file(self.style_guide_path_edit, "JSON 파일 (*.json)"))
        style_guide_layout.addWidget(browse_style_btn)
        paths_layout.addLayout(style_guide_layout)

        # 용어집 경로
        terminology_layout = QHBoxLayout()
        terminology_layout.addWidget(QLabel("용어집:"))
        self.terminology_path_edit = QLineEdit()
        self.terminology_path_edit.setText("config/terminology.json")
        self.terminology_path_edit.setReadOnly(True)
        terminology_layout.addWidget(self.terminology_path_edit)
        browse_term_btn = QPushButton("찾아보기")
        browse_term_btn.clicked.connect(lambda: self.browse_file(self.terminology_path_edit, "JSON 파일 (*.json)"))
        terminology_layout.addWidget(browse_term_btn)
        paths_layout.addLayout(terminology_layout)

        layout.addWidget(paths_group)
        layout.addStretch()

        return tab

    def create_terminology_tab(self):
        """용어집 편집 탭"""
        tab = QWidget()
        layout = QVBoxLayout()
        tab.setLayout(layout)

        # 설명
        info_label = QLabel("📝 용어집을 직접 편집할 수 있습니다. JSON 형식을 유지해주세요.")
        layout.addWidget(info_label)

        # 용어집 편집기
        self.terminology_editor = QTextEdit()
        self.terminology_editor.setPlaceholderText("용어집을 불러오는 중...")
        layout.addWidget(self.terminology_editor)

        # 버튼들
        button_layout = QHBoxLayout()

        load_term_btn = QPushButton("📂 불러오기")
        load_term_btn.clicked.connect(self.load_terminology)
        button_layout.addWidget(load_term_btn)

        save_term_btn = QPushButton("💾 저장")
        save_term_btn.clicked.connect(self.save_terminology)
        button_layout.addWidget(save_term_btn)

        validate_term_btn = QPushButton("✅ 유효성 검사")
        validate_term_btn.clicked.connect(self.validate_terminology)
        button_layout.addWidget(validate_term_btn)

        layout.addLayout(button_layout)

        # 초기 로드
        self.load_terminology()

        return tab

    def create_info_tab(self):
        """시스템 정보 탭"""
        tab = QWidget()
        layout = QVBoxLayout()
        tab.setLayout(layout)

        info_text = QTextEdit()
        info_text.setReadOnly(True)
        info_text.setHtml("""
        <h3>📋 시스템 정보</h3>
        <ul>
            <li><b>버전:</b> 1.0.0</li>
            <li><b>엔진:</b> Claude API (Sonnet 4.5)</li>
            <li><b>지원 입력 형식:</b> .txt, .docx, .pdf</li>
            <li><b>지원 출력 형식:</b> .txt, .docx</li>
        </ul>

        <h3>💡 사용 팁</h3>
        <ul>
            <li>Word → Word 번역을 권장합니다 (실무 최적화)</li>
            <li>자체 검수를 활성화하면 더 높은 품질을 보장합니다</li>
            <li>TM을 활용하면 동일한 문장을 빠르게 번역할 수 있습니다</li>
            <li>자동 섹션 분류로 전체 명세서를 한 번에 번역할 수 있습니다</li>
        </ul>

        <h3>📚 문서</h3>
        <ul>
            <li>README.md - 전체 사용 가이드</li>
            <li>QUICKSTART.md - 빠른 시작</li>
            <li>INPUT_OUTPUT_FORMATS.md - 파일 형식 가이드</li>
            <li>AUTO_SECTION_FEATURE.md - 자동 섹션 분류</li>
            <li>LOGGING_AND_SETTINGS.md - 로깅 및 설정</li>
        </ul>
        """)
        layout.addWidget(info_text)

        return tab

    def select_input_file(self):
        """입력 파일 선택"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "입력 파일 선택",
            "",
            "지원 파일 (*.txt *.docx *.pdf);;모든 파일 (*.*)"
        )
        if file_path:
            self.input_file_edit.setText(file_path)

            # 출력 파일명 자동 생성
            input_path = Path(file_path)
            output_name = f"{input_path.stem}_ko.docx"
            output_path = input_path.parent / output_name
            self.output_file_edit.setText(str(output_path))

    def select_output_file(self):
        """출력 파일 선택"""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "출력 파일 선택",
            "",
            "Word 문서 (*.docx);;텍스트 파일 (*.txt);;모든 파일 (*.*)"
        )
        if file_path:
            self.output_file_edit.setText(file_path)

    def toggle_auto_section(self, checked):
        """자동 섹션 분류 토글"""
        # 자동 분류 모드일 때는 문서 유형 선택 비활성화
        self.doc_type_combo.setEnabled(not checked)

        if checked:
            self.log_text.append(
                "ℹ️ 자동 섹션 분류 모드: 제목, 요약서, 청구항, 명세서를 자동으로 구분하여 번역합니다.\n"
            )
        else:
            self.log_text.append(
                "ℹ️ 일반 모드: 선택한 문서 유형으로 전체 문서를 번역합니다.\n"
            )

    def start_translation(self):
        """번역 시작"""
        input_file = self.input_file_edit.text()
        output_file = self.output_file_edit.text()

        # 유효성 검사
        if not input_file or not output_file:
            QMessageBox.warning(self, "경고", "입력 파일과 출력 파일을 모두 선택해주세요.")
            return

        if not Path(input_file).exists():
            QMessageBox.warning(self, "경고", "입력 파일이 존재하지 않습니다.")
            return

        # 문서 유형
        doc_type_text = self.doc_type_combo.currentText()
        doc_type = doc_type_text.split()[0]  # "청구항", "명세서", "요약서"
        doc_type_map = {"청구항": "claim", "명세서": "specification", "요약서": "abstract"}
        doc_type = doc_type_map.get(doc_type, "claim")

        # 옵션
        use_review = self.review_checkbox.isChecked()
        save_tm = self.tm_checkbox.isChecked()
        auto_section = self.auto_section_checkbox.isChecked()

        # UI 비활성화
        self.translate_btn.setEnabled(False)
        self.progress_bar.show()
        self.log_text.clear()

        if auto_section:
            self.log_text.append("🤖 자동 섹션 분류 모드로 번역을 시작합니다...\n")
        else:
            self.log_text.append("번역을 시작합니다...\n")

        # 번역 스레드 시작
        self.translation_thread = TranslationThread(
            input_file, output_file, doc_type, use_review, save_tm, auto_section
        )
        self.translation_thread.progress.connect(self.update_progress)
        self.translation_thread.finished.connect(self.translation_finished)
        self.translation_thread.error.connect(self.translation_error)
        self.translation_thread.start()

    def update_progress(self, message):
        """진행 상황 업데이트"""
        self.log_text.append(message)
        self.statusBar().showMessage(message)

    def translation_finished(self, result):
        """번역 완료"""
        self.progress_bar.hide()
        self.translate_btn.setEnabled(True)

        translation = result["translation"]
        qa_result = result.get("qa_result", {})

        # 결과 표시
        self.log_text.append("\n" + "="*50)
        self.log_text.append("✅ 번역 완료!")
        self.log_text.append("="*50)
        self.log_text.append(f"\n📄 번역 결과 미리보기:\n")
        self.log_text.append(translation[:500] + "..." if len(translation) > 500 else translation)

        if qa_result:
            passed = qa_result.get("passed", False)
            total_violations = qa_result.get("total_violations", 0)
            self.log_text.append(f"\n📊 QA 결과: {'✅ PASS' if passed else '❌ FAIL'}")
            self.log_text.append(f"   위반 사항: {total_violations}개")

        # 성공 메시지
        QMessageBox.information(
            self,
            "번역 완료",
            f"번역이 완료되었습니다!\n\n출력 파일: {self.output_file_edit.text()}"
        )

        self.statusBar().showMessage("번역 완료")

    def translation_error(self, error_msg):
        """번역 오류"""
        self.progress_bar.hide()
        self.translate_btn.setEnabled(True)

        self.log_text.append(f"\n❌ 오류 발생: {error_msg}")

        QMessageBox.critical(
            self,
            "오류",
            f"번역 중 오류가 발생했습니다:\n\n{error_msg}"
        )

        self.statusBar().showMessage("오류 발생")

    def search_tm(self):
        """TM 검색"""
        query = self.tm_search_input.text().strip()
        if not query:
            self.tm_search_results.setText("검색어를 입력해주세요.")
            return

        try:
            tm = TranslationMemory()
            threshold = self.tm_threshold_spinbox.value() / 100.0

            # TM에서 유사한 문장 검색
            results = tm.search(query, threshold=threshold, limit=10)
            tm.close()

            if not results:
                self.tm_search_results.setText(f"'{query}'에 대한 검색 결과가 없습니다.\n\n유사도 임계값을 낮춰보세요.")
                return

            # 결과 포맷팅
            text = f"🔍 검색어: {query}\n"
            text += f"📊 총 {len(results)}개 결과 (유사도 {int(threshold*100)}% 이상)\n\n"
            text += "=" * 60 + "\n\n"

            for i, result in enumerate(results, 1):
                similarity = result.get('similarity', 0) * 100
                source = result.get('source', '')
                translation = result.get('translation', '')
                domain = result.get('domain', 'unknown')
                doc_type = result.get('document_type', 'unknown')

                text += f"[{i}] 유사도: {similarity:.1f}%\n"
                text += f"도메인: {domain} | 유형: {doc_type}\n\n"
                text += f"원문:\n{source}\n\n"
                text += f"번역:\n{translation}\n\n"
                text += "-" * 60 + "\n\n"

            self.tm_search_results.setText(text)

        except Exception as e:
            self.tm_search_results.setText(f"검색 중 오류 발생:\n{str(e)}")

    def refresh_tm_stats(self):
        """TM 통계 새로고침"""
        try:
            tm = TranslationMemory()
            stats = tm.get_stats()
            tm.close()

            # 통계 텍스트 생성
            text = f"📊 Translation Memory 통계\n\n"
            text += f"총 항목 수: {stats['total']}개\n\n"

            if stats['by_domain']:
                text += "도메인별:\n"
                for domain, count in stats['by_domain'].items():
                    text += f"  - {domain}: {count}개\n"
                text += "\n"

            if stats['by_type']:
                text += "문서 유형별:\n"
                for doc_type, count in stats['by_type'].items():
                    text += f"  - {doc_type}: {count}개\n"

            self.tm_stats_text.setText(text)

        except Exception as e:
            self.tm_stats_text.setText(f"오류: {str(e)}")

    def toggle_api_key_visibility(self):
        """API 키 표시/숨기기"""
        if self.api_key_edit.echoMode() == QLineEdit.EchoMode.Password:
            self.api_key_edit.setEchoMode(QLineEdit.EchoMode.Normal)
        else:
            self.api_key_edit.setEchoMode(QLineEdit.EchoMode.Password)

    def save_api_key(self):
        """API 키 저장"""
        api_key = self.api_key_edit.text().strip()
        if not api_key:
            QMessageBox.warning(self, "경고", "API 키를 입력해주세요.")
            return

        try:
            # .env 파일에 저장
            env_path = Path(".env")
            if env_path.exists():
                with open(env_path, 'r') as f:
                    lines = f.readlines()

                # 기존 API 키 라인 찾기
                found = False
                for i, line in enumerate(lines):
                    if line.startswith("ANTHROPIC_API_KEY="):
                        lines[i] = f"ANTHROPIC_API_KEY={api_key}\n"
                        found = True
                        break

                if not found:
                    lines.append(f"ANTHROPIC_API_KEY={api_key}\n")

                with open(env_path, 'w') as f:
                    f.writelines(lines)
            else:
                with open(env_path, 'w') as f:
                    f.write(f"ANTHROPIC_API_KEY={api_key}\n")

            # 환경 변수에도 설정
            import os
            os.environ["ANTHROPIC_API_KEY"] = api_key

            QMessageBox.information(self, "성공", "API 키가 저장되었습니다.\n다음 번역부터 적용됩니다.")

        except Exception as e:
            QMessageBox.critical(self, "오류", f"API 키 저장 중 오류가 발생했습니다:\n{str(e)}")

    def browse_file(self, line_edit, file_filter):
        """파일 찾아보기"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "파일 선택",
            "",
            file_filter
        )
        if file_path:
            line_edit.setText(file_path)

    def load_terminology(self):
        """용어집 불러오기"""
        try:
            terminology_path = self.terminology_path_edit.text()
            with open(terminology_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # JSON을 보기 좋게 포맷팅
            import json
            data = json.loads(content)
            formatted = json.dumps(data, indent=2, ensure_ascii=False)

            self.terminology_editor.setText(formatted)

        except FileNotFoundError:
            self.terminology_editor.setText("# 용어집 파일을 찾을 수 없습니다.\n# 새로 생성하려면 JSON 형식으로 작성 후 저장하세요.\n\n{}")
        except Exception as e:
            QMessageBox.critical(self, "오류", f"용어집 불러오기 실패:\n{str(e)}")

    def save_terminology(self):
        """용어집 저장"""
        try:
            content = self.terminology_editor.toPlainText()

            # JSON 유효성 검사
            import json
            data = json.loads(content)

            # 저장
            terminology_path = self.terminology_path_edit.text()
            with open(terminology_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)

            QMessageBox.information(self, "성공", "용어집이 저장되었습니다.")

        except json.JSONDecodeError as e:
            QMessageBox.critical(self, "오류", f"JSON 형식이 올바르지 않습니다:\n{str(e)}\n\n유효성 검사를 먼저 실행해보세요.")
        except Exception as e:
            QMessageBox.critical(self, "오류", f"용어집 저장 실패:\n{str(e)}")

    def validate_terminology(self):
        """용어집 유효성 검사"""
        try:
            content = self.terminology_editor.toPlainText()
            import json
            data = json.loads(content)

            # 기본 구조 검사
            if not isinstance(data, dict):
                raise ValueError("용어집은 JSON 객체(dictionary) 형식이어야 합니다.")

            QMessageBox.information(
                self,
                "유효성 검사 통과",
                f"✅ JSON 형식이 올바릅니다.\n\n총 용어 수: {len(data)}개"
            )

        except json.JSONDecodeError as e:
            QMessageBox.critical(
                self,
                "유효성 검사 실패",
                f"❌ JSON 형식 오류:\n\n{str(e)}\n\n줄 {e.lineno}, 열 {e.colno}"
            )
        except Exception as e:
            QMessageBox.critical(
                self,
                "유효성 검사 실패",
                f"❌ 오류:\n\n{str(e)}"
            )

    def refresh_log_list(self):
        """로그 파일 목록 새로고침"""
        self.log_file_list.clear()

        log_dir = Path("logs")
        if not log_dir.exists():
            log_dir.mkdir(exist_ok=True)
            self.log_file_list.addItem("로그 파일이 없습니다.")
            return

        # 로그 파일 목록 가져오기 (최신순)
        log_files = sorted(log_dir.glob("translation_*.log"), key=lambda x: x.stat().st_mtime, reverse=True)

        if not log_files:
            self.log_file_list.addItem("로그 파일이 없습니다.")
            return

        for log_file in log_files:
            # 파일 크기와 수정 시간 표시
            size = log_file.stat().st_size
            size_kb = size / 1024
            mtime = log_file.stat().st_mtime
            from datetime import datetime
            time_str = datetime.fromtimestamp(mtime).strftime("%Y-%m-%d %H:%M:%S")

            item_text = f"{log_file.name} ({size_kb:.1f} KB) - {time_str}"
            item = QListWidgetItem(item_text)
            item.setData(Qt.ItemDataRole.UserRole, str(log_file))  # 파일 경로 저장
            self.log_file_list.addItem(item)

    def load_log_file(self, item):
        """로그 파일 로드"""
        file_path = item.data(Qt.ItemDataRole.UserRole)
        if not file_path or file_path == "로그 파일이 없습니다.":
            return

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # 전체 로그 내용 저장 (필터링용)
            self.current_log_content = content

            # 표시
            self.filter_log_content(self.log_filter_combo.currentText())

        except Exception as e:
            QMessageBox.critical(self, "오류", f"로그 파일 읽기 실패:\n{str(e)}")

    def filter_log_content(self, filter_level):
        """로그 필터링"""
        if not hasattr(self, 'current_log_content') or not self.current_log_content:
            return

        if filter_level == "전체":
            self.log_content_text.setText(self.current_log_content)
        else:
            # 해당 레벨만 표시
            lines = self.current_log_content.split('\n')
            filtered_lines = [line for line in lines if filter_level in line or not any(
                level in line for level in ['INFO', 'WARNING', 'ERROR', 'DEBUG']
            )]
            self.log_content_text.setText('\n'.join(filtered_lines))


def main():
    """애플리케이션 실행"""
    app = QApplication(sys.argv)
    app.setStyle('Fusion')  # 모던한 스타일

    window = PatentTranslatorGUI()
    window.show()

    sys.exit(app.exec())


if __name__ == "__main__":
    main()
