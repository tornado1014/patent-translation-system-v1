import click
from rich.console import Console
from rich.panel import Panel
from rich.syntax import Syntax
from src.graph import app
from docx import Document
from pathlib import Path

console = Console()

@click.command()
@click.argument('input_file', type=click.Path(exists=True))
@click.option('--use-flash-review', is_flag=True, help="Use the faster Gemini Flash model for the review step.")
def run(input_file, use_flash_review):
    """
    Runs the LangGraph-based patent translation system.
    """
    console.print(Panel("🚀 [bold blue]Starting Patent-Mind Translation System[/bold blue]"))

    input_path = Path(input_file)
    if input_path.suffix.lower() == '.docx':
        doc = Document(input_file)
        source_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
    else: # Assume .txt
        with open(input_file, 'r', encoding='utf-8') as f:
            source_text = f.read()

    if not source_text.strip():
        console.print("[bold red]Error: Input file is empty or contains no text.[/bold red]")
        return

    initial_state = {
        "original_text": source_text,
        "document_type": "claim", # Example, can be made dynamic
        "use_flash_review": use_flash_review,
        "messages": []
    }

    console.print(f"📄 Processing file: {input_file}")
    console.print(f"⚡ Using Flash for review: {'Yes' if use_flash_review else 'No'}")
    
    final_state = app.invoke(initial_state)

    translation = final_state.get("draft_translation", "No translation generated.")
    
    console.print(Panel("[bold green]✅ Translation Complete[/bold green]"))
    
    output_path = input_path.with_name(f"{input_path.stem}_translated.docx")
    
    # Save as .docx
    new_doc = Document()
    for line in translation.split('\n'):
        if line.strip():
            new_doc.add_paragraph(line)
    new_doc.save(output_path)
    console.print(f"📄 Translated document saved to: [cyan]{output_path}[/cyan]")

    # Also print to console
    syntax = Syntax(translation, "text", theme="monokai", line_numbers=True)
    console.print(Panel(syntax, title="Final Translation", border_style="green"))

    review = final_state.get("review_result", {})
    if review:
        console.print(Panel(f"Review Passed: {review.get('passed')}\nFeedback: {review.get('feedback')}", title="QA Review", border_style="cyan"))

if __name__ == "__main__":
    run()
