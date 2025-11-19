#!/usr/bin/env python3
"""
특허 번역 자동화 시스템 - 메인 CLI
"""

import sys
from pathlib import Path

# src 디렉토리를 Python 경로에 추가
sys.path.insert(0, str(Path(__file__).parent / "src"))

import click
from rich.console import Console
from rich.panel import Panel
from rich.syntax import Syntax
from pipeline import TranslationPipeline
from tm_manager import TranslationMemory
from rag_guide import StyleGuideRAG

console = Console()


@click.group()
def cli():
    """🚀 특허 번역 자동화 시스템 (Patent Translation Automation System)"""
    pass


@cli.command()
@click.argument('input_file', type=click.Path(exists=True))
@click.option('-o', '--output', type=click.Path(), help='출력 파일 경로')
@click.option('--type', 'document_type', type=click.Choice(['claim', 'specification', 'abstract']),
              default='claim', help='문서 유형')
@click.option('--model', type=click.Choice(['gemini-2.5-flash', 'gemini-2.5-pro', 'gemini-3-pro-preview']),
              default=None, help='번역에 사용할 Gemini 모델 선택')
@click.option('--no-review', is_flag=True, help='자체 검수 생략')
@click.option('--no-tm', is_flag=True, help='TM 저장 생략')
def translate(input_file, output, document_type, model, no_review, no_tm):
    """특허 문서 번역 (지원: .txt, .docx, .pdf)"""

    console.print(Panel.fit("🌟 특허 번역 시작", style="bold blue"))

    input_path = Path(input_file)
    file_ext = input_path.suffix.lower()

    source_text = ""
    if file_ext == '.txt':
        with open(input_file, 'r', encoding='utf-8') as f:
            source_text = f.read()
    elif file_ext == '.docx':
        try:
            from docx import Document
            doc = Document(input_file)
            source_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        except ImportError:
            console.print("❌ python-docx 패키지가 필요합니다: uv add python-docx", style="red")
            sys.exit(1)
    # PDF 지원은 docling이 로컬에 설치되어 있지 않을 수 있으므로 일단 제거
    # elif file_ext == '.pdf': ...
    else:
        console.print(f"❌ 지원하지 않는 파일 형식: {file_ext}. (.txt, .docx만 지원)", style="red")
        sys.exit(1)

    console.print(f"\n📄 입력 파일: {input_file}")
    console.print(f"📋 문서 유형: {document_type}")
    console.print(f"🔍 자체 검수: {'비활성화' if no_review else '활성화'}")
    console.print(f"💾 TM 저장: {'비활성화' if no_tm else '활성화'}")
    if model:
        console.print(f"🤖 사용할 모델: {model}")
    console.print("")

    pipeline = TranslationPipeline()
    
    # 모델 설정 (사용자가 지정한 경우)
    if model:
        pipeline.translator.set_model(model)

    try:
        result = pipeline.translate_document(
            source_text=source_text,
            document_type=document_type,
            use_self_review=not no_review,
            save_to_tm=not no_tm
        )

        if result["success"]:
            translation = result["translation"]
            console.print("\n" + "="*60, style="green")
            console.print("✅ 번역 완료!", style="bold green")
            console.print("="*60 + "\n", style="green")

            syntax = Syntax(translation, "text", theme="monokai", line_numbers=False)
            console.print(Panel(syntax, title="번역 결과", border_style="green"))

            if output:
                output_path = Path(output)
                output_path.parent.mkdir(parents=True, exist_ok=True)
                if output_path.suffix.lower() == '.docx':
                    try:
                        from docx import Document
                        doc = Document()
                        for line in translation.split('\n'):
                            if line.strip():
                                doc.add_paragraph(line)
                        doc.save(str(output_path))
                        console.print(f"\n💾 Word 파일 저장 완료: {output_path}", style="green")
                    except ImportError:
                        console.print("❌ python-docx 패키지가 필요합니다", style="red")
                else:
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(translation)
                    console.print(f"\n💾 텍스트 파일 저장 완료: {output_path}", style="green")

                if "qa_result" in result:
                    qa_report = pipeline.qa_checker.generate_report(result["qa_result"])
                    qa_path = output_path.with_suffix('.qa.txt')
                    with open(qa_path, 'w', encoding='utf-8') as f:
                        f.write(qa_report)
                    console.print(f"📊 QA 리포트 저장: {qa_path}", style="cyan")

            if "qa_result" in result:
                qa = result["qa_result"]
                console.print(f"\n📊 QA 결과: {'✅ PASS' if qa['passed'] else '❌ FAIL'}")
                console.print(f"   위반 사항: {qa['total_violations']}개")
        else:
            console.print(f"\n❌ 번역 실패: {result.get('error')}", style="bold red")
            sys.exit(1)
    finally:
        pipeline.close()


@cli.command()
def tm_stats():
    """Translation Memory 통계"""
    tm = TranslationMemory()
    try:
        stats = tm.get_stats()
        console.print(Panel.fit("📊 Translation Memory 통계", style="bold cyan"))
        console.print(f"\n총 항목 수: {stats['total']}개\n")
        if stats['by_domain']:
            console.print("도메인별:")
            for domain, count in stats['by_domain'].items():
                console.print(f"  - {domain}: {count}개")
        if stats['by_type']:
            console.print("\n문서 유형별:")
            for doc_type, count in stats['by_type'].items():
                console.print(f"  - {doc_type}: {count}개")
    finally:
        tm.close()


@cli.command()
@click.argument('guide_path', type=click.Path(exists=True))
def init_rag(guide_path):
    """스타일 가이드 RAG 인덱싱"""
    console.print(Panel.fit("🔧 스타일 가이드 인덱싱", style="bold yellow"))
    console.print(f"\n📄 파일: {guide_path}")
    rag = StyleGuideRAG()
    if rag.index_style_guide(guide_path):
        console.print("\n✅ 인덱싱 완료!", style="green")
    else:
        console.print("\n❌ 인덱싱 실패", style="red")
        sys.exit(1)


@cli.command()
def version():
    """버전 정보"""
    console.print(Panel.fit(
        """
🚀 특허 번역 자동화 시스템
버전: 1.1.0 (Gemini API)
Python + Google Gemini 기반
        """,
        style="bold magenta"
    ))


if __name__ == "__main__":
    cli()
